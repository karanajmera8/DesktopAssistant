import pyttsx3
import webbrowser
import wikipedia
import speech_recognition as sr
import datetime
import openai
import docx
import pvporcupine
import pyaudio
import struct
import time
import win32com.client as win32
from slack_sdk import WebClient
import os
from pyOutlook import OutlookAccount

def takeCommand():
    r=sr.Recognizer()

    with sr.Microphone() as source:
        print('Listening')
        r.pause_threshold = 0.7
        audio = r.listen(source)

        try:
            print('Recognizing')

            Query = r.recognize_google(audio,language = 'en-in')
            print("The Command is Printed = ",Query)

        except Exception as e:
            print(e)
            speak("Can you repeat that please")
            return takeCommand().lower()

        return Query

def speak(audio):

    engine = pyttsx3.init()
    voices = engine.getProperty('voices')

    engine.setProperty('voice',voices[0].id)

    engine.say(audio)

    engine.runAndWait()

def tellDay():
    day = datetime.datetime.today().weekday() +1

    day_dict={
    1:'Monday',
    2:'Tuesday',
    3:'Wednesday',
    4:'Thursday',
    5:'Friday',
    6:'Saturday',
    7:'Sunday'
    }

    if day in day_dict.keys():
        print(day_dict[day])
        speak("The day is " + day_dict[day])

def tellTime():
    time = str(datetime.datetime.now())

    print(time)
    hour = time[11:13]
    min = time[14:16]
    speak("The time is sir"+hour+"Hours and"+min+"Minutes")

def Hello():
    speak("Hello sir I am your desktop assistant. Tell me how may I help you")

def chatgpt():
    openai.api_key = "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"
    messages=[
    {'role':'system','content':'You are an intelligent assistant.'}
    ]
    speak("Welcome to chat GPT!")
    while True:
        message = takeCommand().lower()
        if "exit" in message:
            return
        if message:
            messages.append({'role':'user','content':message})
            chat = openai.chat.completions.create(model="gpt-4",messages=messages)
            reply=chat.choices[0].message.content
            speak(reply)
            messages.append({'role':'assistant','content':reply})
def word_file():
    doc=docx.Document()
    p = doc.add_paragraph()
    speak("Please start saying your content")
    run=p.add_run(takeCommand())
    doc.save("DesktopAssistan_demo.docx")
    speak("content saved in a word document")

def youtube():
    speak("What video do you wish to play")
    video_play=takeCommand()
    webbrowser.open("www.youtube.com/"+video_play)

def sendemail():
    mailbox = win32.Dispatch('outlook.application')
    email = mailbox.CreateItem(0)
    speak("what email provider do you want to send the email to")
    email_provider = takeCommand().lower()
    speak("Who do you want to send the email to")
    email.To = takeCommand().lower()+'@'+email_provider+'.com'
    speak("What should be the subject of the email")
    email.Subject = takeCommand().lower()
    speak("WHat is the content of the email")
    email.Body = takeCommand().lower()

    email.Send()

def slack():
    client = WebClient(token=os.environ.get("SLACK_BOT_TOKEN"))

    result = client.chat_postMessage(channel="D02MY8PF17D",text="Hello")

def wake_detect():
    speak("How can I help you")
    while(True):

        query = takeCommand().lower()

        if "open geeksforgeeks" in query:
            speak("Opening GeeksforGeeks")

            webbrowser.open("www.geeksforgeeks.com")
            continue

        elif "open google" in query:
            speak("Opening Google")

            webbrowser.open("www.google.com")
            continue

        elif "which day it is" in query:
            tellDay()
            continue

        elif "tell me the time" in query:
            tellTime()
            continue

        elif "bye" in query:
            speak("Goodbye")
            break

        elif "tell me your name" in query:
            speak("I am Jarvis. Your desktop Assistant")

        elif "chat gpt" in query:
            chatgpt()

        elif "document" in query:
            word_file()

        elif "youtube" in query:
            youtube()

        elif "send email" in query:
            sendemail()

        elif "slack" in query:
            slack()

        else:
            speak("Sorry I am not trained for that yet")
        speak("What else can I help you with")

def Take_query():

    speak("Hello, this is your virtual assistant Jarvis! Please call out for me for anything I can help you with")
    while True:
        porcupine = None
        pa=None
        audio_stream = None

        try:
            porcupine=pvporcupine.create(keywords=["computer"])
            pa = pyaudio.PyAudio()
            audio_stream = pa.open(
                            rate=porcupine.sample_rate,
                            channels=1,
                            format=pyaudio.paInt16,
                            input=True,
                            frames_per_buffer=porcupine.frame_length)
            while True:
                pcm = audio_stream.read(porcupine.frame_length)
                pcm = struct.unpack_from("h" * porcupine.frame_length,pcm)
                keyword_index = porcupine.process(pcm)
                if keyword_index >= 0:
                    wake_detect()
                    time.sleep(1)
        finally:
            if porcupine is not None:
                porcupine.delete()

            if audio_stream is not None:
                audio_stream.close()

            if pa is not None:
                pa.terminate()

Take_query()
